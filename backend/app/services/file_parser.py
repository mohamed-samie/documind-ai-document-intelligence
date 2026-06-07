import csv
import re
from pathlib import Path
from typing import Dict, List, Optional

import fitz
from docx import Document
from openpyxl import load_workbook

SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.xlsx', '.xlsm', '.csv'}
_TEMPLATE_PATTERNS = [
    re.compile(r'power\s*point\s*template', re.I),
    re.compile(r'whirlwind', re.I),
    re.compile(r'example@example\.com', re.I),
    re.compile(r'www\.example\.com', re.I),
    re.compile(r'this is a sample text', re.I),
]


def _clean_lines(text: str) -> List[str]:
    lines = []
    for raw in (text or '').replace('\r', '\n').split('\n'):
        line = re.sub(r'\s+', ' ', raw).strip()
        if not line:
            continue
        if any(pattern.search(line) for pattern in _TEMPLATE_PATTERNS):
            continue
        lines.append(line)
    return lines


def _clean_text(text: str) -> str:
    return '\n'.join(_clean_lines(text))


def _looks_like_heading(text: str) -> bool:
    text = re.sub(r'\s+', ' ', (text or '').strip())
    if len(text) < 2 or len(text) > 100:
        return False
    if any(pattern.search(text) for pattern in _TEMPLATE_PATTERNS):
        return False
    if text.endswith('.') and len(text.split()) > 5:
        return False
    words = text.split()
    if len(words) > 14:
        return False
    if len(text) <= 4 and text.isalnum():
        return True
    title_words = sum(1 for word in words if word[:1].isupper() or word.isupper())
    return title_words >= max(1, len(words) // 2)


def _normalise_title(title: Optional[str]) -> str:
    if not title:
        return ''
    return re.sub(r'\s+', ' ', title).strip()


def _extract_pdf_blocks(page: fitz.Page) -> List[Dict]:
    blocks: List[Dict] = []
    for block in page.get_text('dict').get('blocks', []):
        if block.get('type') != 0:
            continue
        text_parts: List[str] = []
        max_font_size = 0.0
        for line in block.get('lines', []):
            line_parts = []
            for span in line.get('spans', []):
                value = span.get('text', '')
                if value:
                    line_parts.append(value)
                max_font_size = max(max_font_size, float(span.get('size', 0.0)))
            if line_parts:
                text_parts.append(' '.join(line_parts))
        text = re.sub(r'\s+', ' ', ' '.join(text_parts)).strip()
        if text and not any(pattern.search(text) for pattern in _TEMPLATE_PATTERNS):
            blocks.append({'text': text, 'font_size': max_font_size, 'y0': float(block['bbox'][1])})
    blocks.sort(key=lambda item: item['y0'])
    return blocks


def _detect_pdf_title(page: fitz.Page, text: str) -> str:
    blocks = _extract_pdf_blocks(page)
    candidates: List[str] = []
    for block in sorted(blocks, key=lambda item: -item['font_size'])[:6]:
        candidates.append(block['text'])
    lines = _clean_lines(text)
    candidates.extend(lines[:6])
    candidates.extend(lines[-4:])

    for candidate in candidates:
        candidate = _normalise_title(candidate)
        if _looks_like_heading(candidate):
            return candidate
    return ''


def parse_pdf(file_path: str) -> List[Dict]:
    pages: List[Dict] = []
    document = fitz.open(file_path)
    current_section_index = 0
    current_section_title = ''
    previous_title = ''

    try:
        for page_index, page in enumerate(document):
            raw_text = page.get_text().strip()
            text = _clean_text(raw_text)
            title = _detect_pdf_title(page, raw_text)

            if title and title.lower() != previous_title.lower():
                current_section_index += 1
                current_section_title = title
                previous_title = title
            elif current_section_index == 0:
                current_section_index = 1
                current_section_title = title or Path(file_path).stem

            pages.append({
                'page_number': page_index + 1,
                'text': text,
                'title': title,
                'section_index': current_section_index,
                'section_title': current_section_title,
            })
    finally:
        document.close()
    return [page for page in pages if page.get('text')]


def parse_docx(file_path: str) -> List[Dict]:
    doc = Document(file_path)
    pages: List[Dict] = []
    current_title = Path(file_path).stem
    current_section_index = 1
    current_lines: List[str] = []

    def flush():
        nonlocal current_lines
        text = _clean_text('\n'.join(current_lines))
        if text:
            pages.append({
                'page_number': len(pages) + 1,
                'text': text,
                'title': current_title,
                'section_index': current_section_index,
                'section_title': current_title,
            })
        current_lines = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name or '').lower()
        if 'heading' in style:
            flush()
            current_section_index += 1
            current_title = text
            current_lines = [text]
        else:
            current_lines.append(text)
    flush()

    if not pages:
        text = _clean_text('\n'.join(p.text for p in doc.paragraphs if p.text.strip()))
        if text:
            pages.append({'page_number': 1, 'text': text, 'title': current_title, 'section_index': 1, 'section_title': current_title})
    return pages


def parse_xlsx(file_path: str) -> List[Dict]:
    workbook = load_workbook(file_path, read_only=True, data_only=True)
    pages: List[Dict] = []
    try:
        for sheet_index, sheet in enumerate(workbook.worksheets, start=1):
            rows: List[str] = []
            for row in sheet.iter_rows(values_only=True):
                values = [str(value).strip() for value in row if value is not None and str(value).strip()]
                if values:
                    rows.append(' | '.join(values))
            text = _clean_text('\n'.join(rows))
            if text:
                pages.append({'page_number': sheet_index, 'text': text, 'title': sheet.title, 'section_index': sheet_index, 'section_title': sheet.title})
    finally:
        workbook.close()
    return pages


def parse_csv(file_path: str) -> List[Dict]:
    rows: List[str] = []
    with open(file_path, 'r', encoding='utf-8-sig', newline='') as handle:
        reader = csv.reader(handle)
        for row in reader:
            values = [cell.strip() for cell in row if cell and cell.strip()]
            if values:
                rows.append(' | '.join(values))
    text = _clean_text('\n'.join(rows))
    title = Path(file_path).stem
    return [{'page_number': 1, 'text': text, 'title': title, 'section_index': 1, 'section_title': title}] if text else []


def parse_file(file_path: str) -> List[Dict]:
    extension = Path(file_path).suffix.lower()
    if extension == '.pdf':
        return parse_pdf(file_path)
    if extension == '.docx':
        return parse_docx(file_path)
    if extension in {'.xlsx', '.xlsm'}:
        return parse_xlsx(file_path)
    if extension == '.csv':
        return parse_csv(file_path)
    raise ValueError(f"Unsupported file type: {extension}. Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")
